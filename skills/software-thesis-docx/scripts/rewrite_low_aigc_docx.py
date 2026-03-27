from __future__ import annotations

import argparse
import json
import shutil
from pathlib import Path
from typing import Any

from docx import Document

from aigc_utils import parse_paragraph_indices, rewrite_text_by_recipe


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Rewrite authorized DOCX paragraphs for lower AIGC-style risk while preserving layout where possible."
    )
    parser.add_argument("--input", required=True, help="Input DOCX path.")
    parser.add_argument("--report", required=True, help="Risk report JSON produced by check_aigc_risk.py.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    parser.add_argument("--pending-output", required=True, help="Output JSON path for paragraphs that need manual review.")
    parser.add_argument(
        "--profile",
        choices=["academic_safe", "explicit_low_aigc"],
        default="academic_safe",
        help="Rewrite profile. Use explicit_low_aigc only when the user explicitly asks to lower AIGC.",
    )
    parser.add_argument(
        "--authorized-indices",
        help="Optional comma-separated 0-based paragraph indices. If omitted, the script rewrites report items with recommended_rewrite=true.",
    )
    parser.add_argument(
        "--normalize-typography",
        action="store_true",
        help="Apply punctuation and spacing normalization to targeted paragraphs.",
    )
    return parser.parse_args()


def load_report(path: Path) -> dict[int, dict[str, Any]]:
    with path.open("r", encoding="utf-8") as fh:
        data = json.load(fh)
    items = data.get("items")
    if not isinstance(items, list):
        raise ValueError("Report JSON must contain an 'items' array.")
    result: dict[int, dict[str, Any]] = {}
    for item in items:
        index = item.get("paragraph_index")
        if not isinstance(index, int):
            continue
        result[index] = item
    return result


def target_indices(report_items: dict[int, dict[str, Any]], authorized: list[int] | None) -> list[int]:
    if authorized is not None:
        return authorized
    return sorted(index for index, item in report_items.items() if bool(item.get("recommended_rewrite")))


def replace_single_run(paragraph, new_text: str) -> None:
    if len(paragraph.runs) != 1:
        raise ValueError(f"Target paragraph is not single-run: {len(paragraph.runs)}")
    paragraph.runs[0].text = new_text


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).resolve()
    report_path = Path(args.report).resolve()
    output_path = Path(args.output).resolve()
    pending_output_path = Path(args.pending_output).resolve()
    authorized = parse_paragraph_indices(args.authorized_indices)

    report_items = load_report(report_path)
    indices = target_indices(report_items, authorized)

    shutil.copy2(input_path, output_path)
    document = Document(output_path)

    applied_items: list[dict[str, Any]] = []
    pending_items: list[dict[str, Any]] = []

    for index in indices:
        if index < 0 or index >= len(document.paragraphs):
            pending_items.append(
                {
                    "paragraph_index": index,
                    "reason": "paragraph_index_out_of_range",
                    "suggested_text": None,
                }
            )
            continue

        paragraph = document.paragraphs[index]
        report_item = report_items.get(index, {})
        recipes = report_item.get("rewrite_recipe") or []
        original_text = paragraph.text
        suggested_text = rewrite_text_by_recipe(
            original_text,
            recipes,
            args.profile,
            normalize_typography_enabled=bool(args.normalize_typography),
        )

        if len(paragraph.runs) != 1:
            pending_items.append(
                {
                    "paragraph_index": index,
                    "reason": "multi_run_paragraph",
                    "run_count": len(paragraph.runs),
                    "excerpt": original_text[:160],
                    "rewrite_recipe": recipes,
                    "suggested_text": suggested_text,
                }
            )
            continue

        if suggested_text == original_text:
            pending_items.append(
                {
                    "paragraph_index": index,
                    "reason": "no_change_after_rewrite_recipe",
                    "run_count": len(paragraph.runs),
                    "excerpt": original_text[:160],
                    "rewrite_recipe": recipes,
                    "suggested_text": suggested_text,
                }
            )
            continue

        replace_single_run(paragraph, suggested_text)
        applied_items.append(
            {
                "paragraph_index": index,
                "rewrite_recipe": recipes,
                "excerpt": original_text[:160],
                "rewritten_excerpt": suggested_text[:160],
            }
        )

    document.save(output_path)

    pending_output_path.parent.mkdir(parents=True, exist_ok=True)
    with pending_output_path.open("w", encoding="utf-8") as fh:
        json.dump(
            {
                "schema_version": 1,
                "profile": args.profile,
                "normalize_typography": bool(args.normalize_typography),
                "requested_indices": indices,
                "applied_count": len(applied_items),
                "pending_count": len(pending_items),
                "applied_items": applied_items,
                "pending_items": pending_items,
            },
            fh,
            ensure_ascii=False,
            indent=2,
        )
        fh.write("\n")

    print(f"Wrote {output_path}")
    print(f"Wrote {pending_output_path}")


if __name__ == "__main__":
    main()

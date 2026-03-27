from __future__ import annotations

import copy
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ALIGNMENT_NAMES = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
}

LINE_SPACING_RULE_NAMES = {
    WD_LINE_SPACING.SINGLE: "single",
    WD_LINE_SPACING.EXACTLY: "exactly",
    WD_LINE_SPACING.MULTIPLE: "multiple",
}

ALIGNMENT_VALUES = {value: key for key, value in ALIGNMENT_NAMES.items()}
LINE_SPACING_RULE_VALUES = {value: key for key, value in LINE_SPACING_RULE_NAMES.items()}

DEFAULT_PRESET_FILENAME = "default_style_preset.json"


def default_style_preset_path() -> Path:
    return Path(__file__).resolve().parent.parent / "assets" / "presets" / DEFAULT_PRESET_FILENAME


def length_to_cm(value) -> float | None:
    if value is None:
        return None
    return round(value.cm, 4)


def length_to_pt(value) -> float | None:
    if value is None:
        return None
    return round(value.pt, 2)


def pt_to_length(value: float | None):
    return Pt(value) if value is not None else None


def cm_to_length(value: float | None):
    return Cm(value) if value is not None else None


def normalize_font_name(value: str | None, fallback: str) -> str:
    return value if value else fallback


def run_fonts(run) -> dict[str, str | None]:
    ascii_font = run.font.name
    east_asia_font = None
    rpr = run._element.rPr
    if rpr is not None and rpr.rFonts is not None:
        east_asia_font = rpr.rFonts.get(qn("w:eastAsia"))
        ascii_font = rpr.rFonts.get(qn("w:ascii")) or ascii_font
    return {
        "ascii_font": normalize_font_name(ascii_font, "Times New Roman"),
        "east_asia_font": normalize_font_name(east_asia_font, ascii_font or "宋体"),
    }


def extract_run_spec(run) -> dict[str, Any]:
    fonts = run_fonts(run)
    return {
        "ascii_font": fonts["ascii_font"],
        "east_asia_font": fonts["east_asia_font"],
        "size_pt": length_to_pt(run.font.size) or 12.0,
        "bold": bool(run.font.bold),
        "italic": bool(run.font.italic),
    }


def extract_paragraph_spec(paragraph) -> dict[str, Any]:
    pf = paragraph.paragraph_format
    alignment = ALIGNMENT_NAMES.get(paragraph.alignment, "left")
    line_rule = LINE_SPACING_RULE_NAMES.get(pf.line_spacing_rule, "single")
    line_value = pf.line_spacing
    if hasattr(line_value, "pt"):
        line_value = length_to_pt(line_value)
    elif isinstance(line_value, (int, float)):
        line_value = round(float(line_value), 3)
    else:
        line_value = None
    return {
        "alignment": alignment,
        "first_line_indent_cm": length_to_cm(pf.first_line_indent),
        "left_indent_cm": length_to_cm(pf.left_indent),
        "space_before_pt": length_to_pt(pf.space_before),
        "space_after_pt": length_to_pt(pf.space_after),
        "line_spacing_rule": line_rule,
        "line_spacing_value": line_value,
    }


def role_spec_from_paragraph(paragraph, run_index: int = 0) -> dict[str, Any]:
    run = paragraph.runs[run_index] if paragraph.runs else None
    if run is None:
        raise ValueError(f"Paragraph has no runs: {paragraph.text[:60]}")
    return {
        "paragraph": extract_paragraph_spec(paragraph),
        "run": extract_run_spec(run),
    }


def legacy_style_preset() -> dict[str, Any]:
    return {
        "schema_version": 1,
        "source": {
            "kind": "built_in_fallback",
            "label": "software-thesis-docx legacy preset",
        },
        "page": {
            "top_margin_cm": 2.5,
            "bottom_margin_cm": 2.5,
            "left_margin_cm": 2.0,
            "right_margin_cm": 2.0,
            "header_distance_cm": 1.5,
            "footer_distance_cm": 1.75,
        },
        "roles": {
            "cover_label": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 6.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 24.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "黑体",
                    "size_pt": 20.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "cover_title_cn": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 10.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 24.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "黑体",
                    "size_pt": 22.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "cover_title_en": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 20.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 24.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "Times New Roman",
                    "size_pt": 14.0,
                    "bold": False,
                    "italic": False,
                },
            },
            "cover_meta": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 4.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 24.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "宋体",
                    "size_pt": 14.0,
                    "bold": False,
                    "italic": False,
                },
            },
            "abstract_heading_cn": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 12.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 24.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "黑体",
                    "size_pt": 16.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "abstract_heading_en": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 12.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 24.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "Times New Roman",
                    "size_pt": 16.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "body_cn": {
                "paragraph": {
                    "alignment": "justify",
                    "first_line_indent_cm": 0.85,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 0.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 20.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "宋体",
                    "size_pt": 12.0,
                    "bold": False,
                    "italic": False,
                },
            },
            "body_en": {
                "paragraph": {
                    "alignment": "left",
                    "first_line_indent_cm": 0.85,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 0.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 20.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "Times New Roman",
                    "size_pt": 12.0,
                    "bold": False,
                    "italic": False,
                },
            },
            "keywords_label_cn": {
                "paragraph": {
                    "alignment": "left",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 0.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 20.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "黑体",
                    "size_pt": 12.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "keywords_text_cn": {
                "paragraph": {
                    "alignment": "left",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 0.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 20.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "宋体",
                    "size_pt": 12.0,
                    "bold": False,
                    "italic": False,
                },
            },
            "keywords_label_en": {
                "paragraph": {
                    "alignment": "left",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 0.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 20.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "Times New Roman",
                    "size_pt": 12.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "keywords_text_en": {
                "paragraph": {
                    "alignment": "left",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 0.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 20.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "Times New Roman",
                    "size_pt": 12.0,
                    "bold": False,
                    "italic": False,
                },
            },
            "toc_heading": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 12.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 24.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "黑体",
                    "size_pt": 16.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "chapter": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 12.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 24.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "黑体",
                    "size_pt": 16.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "section": {
                "paragraph": {
                    "alignment": "left",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 6.0,
                    "space_after_pt": 6.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 22.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "黑体",
                    "size_pt": 14.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "subsection": {
                "paragraph": {
                    "alignment": "left",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 3.0,
                    "space_after_pt": 3.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 20.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "黑体",
                    "size_pt": 12.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "table_caption": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 6.0,
                    "space_after_pt": 4.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 18.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "宋体",
                    "size_pt": 10.5,
                    "bold": True,
                    "italic": False,
                },
            },
            "figure_caption": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 3.0,
                    "space_after_pt": 6.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 18.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "宋体",
                    "size_pt": 10.5,
                    "bold": False,
                    "italic": False,
                },
            },
            "reference_heading": {
                "paragraph": {
                    "alignment": "center",
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm": 0.0,
                    "space_before_pt": 0.0,
                    "space_after_pt": 12.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 24.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "黑体",
                    "size_pt": 16.0,
                    "bold": True,
                    "italic": False,
                },
            },
            "reference_item": {
                "paragraph": {
                    "alignment": "left",
                    "first_line_indent_cm": -0.74,
                    "left_indent_cm": 0.74,
                    "space_before_pt": 0.0,
                    "space_after_pt": 0.0,
                    "line_spacing_rule": "exactly",
                    "line_spacing_value": 18.0,
                },
                "run": {
                    "ascii_font": "Times New Roman",
                    "east_asia_font": "宋体",
                    "size_pt": 10.5,
                    "bold": False,
                    "italic": False,
                },
            },
        },
    }


SEMANTIC_PATTERNS = {
    "abstract_heading_cn": re.compile(r"^摘要$"),
    "abstract_heading_en": re.compile(r"^Abstract$", re.IGNORECASE),
    "keywords_cn": re.compile(r"^关键词[:：]"),
    "keywords_en": re.compile(r"^(Keywords|Key words)[:：]", re.IGNORECASE),
    "chapter": re.compile(r"^第[一二三四五六七八九十百]+章"),
    "subsection": re.compile(r"^\d+\.\d+\.\d+"),
    "section": re.compile(r"^\d+\.\d+(?!\.)"),
    "references_heading": re.compile(r"^参考文献$"),
    "figure_caption": re.compile(r"^图\s*\d+[-－.]\d+"),
    "table_caption": re.compile(r"^表\s*\d+[-－.]\d+"),
}


def first_match(document: Document, pattern: re.Pattern[str], *, start_index: int = 0):
    for index, paragraph in enumerate(document.paragraphs[start_index:], start=start_index):
        if pattern.match(paragraph.text.strip()):
            return index, paragraph
    return None, None


def first_long_paragraph(document: Document, *, start_index: int = 0, ascii_only: bool | None = None):
    for index, paragraph in enumerate(document.paragraphs[start_index:], start=start_index):
        text = paragraph.text.strip()
        if len(text) < 40:
            continue
        if ascii_only is True and any("\u4e00" <= ch <= "\u9fff" for ch in text):
            continue
        if ascii_only is False and not any("\u4e00" <= ch <= "\u9fff" for ch in text):
            continue
        return index, paragraph
    return None, None


def extract_style_preset_from_document(document: Document, source_label: str | None = None) -> dict[str, Any]:
    preset = legacy_style_preset()
    if source_label:
        preset["source"] = {
            "kind": "extracted_docx",
            "label": source_label,
        }

    section = document.sections[0]
    preset["page"] = {
        "top_margin_cm": length_to_cm(section.top_margin) or preset["page"]["top_margin_cm"],
        "bottom_margin_cm": length_to_cm(section.bottom_margin) or preset["page"]["bottom_margin_cm"],
        "left_margin_cm": length_to_cm(section.left_margin) or preset["page"]["left_margin_cm"],
        "right_margin_cm": length_to_cm(section.right_margin) or preset["page"]["right_margin_cm"],
        "header_distance_cm": length_to_cm(section.header_distance) or preset["page"]["header_distance_cm"],
        "footer_distance_cm": length_to_cm(section.footer_distance) or preset["page"]["footer_distance_cm"],
    }

    abstract_cn_idx, abstract_cn = first_match(document, SEMANTIC_PATTERNS["abstract_heading_cn"])
    if abstract_cn is not None and abstract_cn_idx is not None:
        preset["roles"]["abstract_heading_cn"] = role_spec_from_paragraph(abstract_cn)
        _, body_cn = first_long_paragraph(document, start_index=abstract_cn_idx + 1, ascii_only=False)
        if body_cn is not None:
            preset["roles"]["body_cn"] = role_spec_from_paragraph(body_cn)

    abstract_en_idx, abstract_en = first_match(document, SEMANTIC_PATTERNS["abstract_heading_en"])
    if abstract_en is not None and abstract_en_idx is not None:
        preset["roles"]["abstract_heading_en"] = role_spec_from_paragraph(abstract_en)
        _, body_en = first_long_paragraph(document, start_index=abstract_en_idx + 1, ascii_only=True)
        if body_en is not None:
            preset["roles"]["body_en"] = role_spec_from_paragraph(body_en)

    _, keywords_cn = first_match(document, SEMANTIC_PATTERNS["keywords_cn"])
    if keywords_cn is not None and len(keywords_cn.runs) >= 2:
        preset["roles"]["keywords_label_cn"] = {
            "paragraph": extract_paragraph_spec(keywords_cn),
            "run": extract_run_spec(keywords_cn.runs[0]),
        }
        preset["roles"]["keywords_text_cn"] = {
            "paragraph": extract_paragraph_spec(keywords_cn),
            "run": extract_run_spec(keywords_cn.runs[1]),
        }

    _, keywords_en = first_match(document, SEMANTIC_PATTERNS["keywords_en"])
    if keywords_en is not None and len(keywords_en.runs) >= 2:
        preset["roles"]["keywords_label_en"] = {
            "paragraph": extract_paragraph_spec(keywords_en),
            "run": extract_run_spec(keywords_en.runs[0]),
        }
        preset["roles"]["keywords_text_en"] = {
            "paragraph": extract_paragraph_spec(keywords_en),
            "run": extract_run_spec(keywords_en.runs[1]),
        }

    for role_name, preset_role in [
        ("chapter", "chapter"),
        ("section", "section"),
        ("subsection", "subsection"),
        ("figure_caption", "figure_caption"),
        ("table_caption", "table_caption"),
        ("references_heading", "reference_heading"),
    ]:
        _, paragraph = first_match(document, SEMANTIC_PATTERNS[role_name])
        if paragraph is not None:
            preset["roles"][preset_role] = role_spec_from_paragraph(paragraph)

    ref_idx, _ = first_match(document, SEMANTIC_PATTERNS["references_heading"])
    if ref_idx is not None:
        for paragraph in document.paragraphs[ref_idx + 1 :]:
            if paragraph.text.strip():
                preset["roles"]["reference_item"] = role_spec_from_paragraph(paragraph)
                break

    return preset


def load_style_preset(path: Path) -> dict[str, Any]:
    import json

    with path.open("r", encoding="utf-8") as fh:
        data = json.load(fh)
    if not isinstance(data, dict):
        raise ValueError("Style preset root must be a JSON object.")
    return data


def save_style_preset(path: Path, preset: dict[str, Any]) -> None:
    import json

    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as fh:
        json.dump(preset, fh, ensure_ascii=False, indent=2)
        fh.write("\n")


def resolve_role(preset: dict[str, Any], role_name: str) -> dict[str, Any]:
    roles = preset.get("roles", {})
    if role_name not in roles:
        raise KeyError(f"Style preset is missing role '{role_name}'.")
    return roles[role_name]


def apply_run_style(run, run_spec: dict[str, Any]) -> None:
    run.font.name = run_spec["ascii_font"]
    run.font.size = Pt(float(run_spec["size_pt"]))
    run.font.bold = bool(run_spec.get("bold", False))
    run.font.italic = bool(run_spec.get("italic", False))
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        from docx.oxml import OxmlElement

        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), run_spec["ascii_font"])
    rfonts.set(qn("w:hAnsi"), run_spec["ascii_font"])
    rfonts.set(qn("w:eastAsia"), run_spec["east_asia_font"])
    rfonts.set(qn("w:cs"), run_spec["ascii_font"])


def apply_paragraph_style(paragraph, paragraph_spec: dict[str, Any]) -> None:
    paragraph.alignment = ALIGNMENT_VALUES.get(paragraph_spec.get("alignment", "left"), WD_ALIGN_PARAGRAPH.LEFT)
    pf = paragraph.paragraph_format
    pf.first_line_indent = cm_to_length(paragraph_spec.get("first_line_indent_cm"))
    pf.left_indent = cm_to_length(paragraph_spec.get("left_indent_cm"))
    pf.space_before = pt_to_length(paragraph_spec.get("space_before_pt"))
    pf.space_after = pt_to_length(paragraph_spec.get("space_after_pt"))
    line_rule_name = paragraph_spec.get("line_spacing_rule", "single")
    pf.line_spacing_rule = LINE_SPACING_RULE_VALUES.get(line_rule_name, WD_LINE_SPACING.SINGLE)
    line_value = paragraph_spec.get("line_spacing_value")
    if line_value is not None:
        if line_rule_name in {"single", "multiple"}:
            pf.line_spacing = float(line_value)
        else:
            pf.line_spacing = Pt(float(line_value))


def clone_preset(preset: dict[str, Any]) -> dict[str, Any]:
    return copy.deepcopy(preset)

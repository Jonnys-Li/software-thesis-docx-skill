from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

try:
    from PIL import Image
except Exception:
    Image = None


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build a thesis DOCX from a structured manifest.")
    parser.add_argument("--manifest", required=True, help="Path to the thesis manifest JSON file.")
    parser.add_argument("--output", required=True, help="Path to the output DOCX file.")
    parser.add_argument(
        "--workspace",
        help="Optional workspace root used to resolve relative figure paths after manifest-relative resolution.",
    )
    return parser.parse_args()


def add_update_fields_on_open(document: Document) -> None:
    settings = document.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def add_field(paragraph, instruction: str, placeholder: str = ""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)
    return run


def set_run_fonts(run, east_asia_font: str = "宋体", ascii_font: str = "Times New Roman", size: float = 12, bold=None) -> None:
    run.font.name = ascii_font
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    rfonts.set(qn("w:cs"), ascii_font)


def set_style_fonts(style, east_asia_font: str = "宋体", ascii_font: str = "Times New Roman", size: float = 12, bold=None) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size)
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia_font)
    rfonts.set(qn("w:cs"), ascii_font)


def ensure_custom_style(document: Document, name: str) -> None:
    if name in document.styles:
        return
    document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def fit_image(path: Path, max_width_cm: float, max_height_cm: float):
    if Image is None:
        return Cm(max_width_cm), None
    with Image.open(path) as img:
        width_px, height_px = img.size
    width_cm = width_px / 96 * 2.54
    height_cm = height_px / 96 * 2.54
    scale = min(max_width_cm / width_cm, max_height_cm / height_cm, 1.0)
    return Cm(width_cm * scale), Cm(height_cm * scale)


class ThesisDocBuilder:
    def __init__(self, manifest: dict[str, Any], manifest_path: Path, output_path: Path, workspace: Path | None) -> None:
        self.manifest = manifest
        self.manifest_path = manifest_path
        self.output_path = output_path
        self.workspace = workspace
        self.document = Document()
        self._configure_styles()
        self._configure_sections()
        add_update_fields_on_open(self.document)

    def _configure_styles(self) -> None:
        ensure_custom_style(self.document, "Caption")
        style_config = [
            ("Normal", "宋体", 12, False),
            ("Heading 1", "黑体", 16, True),
            ("Heading 2", "黑体", 14, True),
            ("Heading 3", "黑体", 12, True),
            ("Title", "黑体", 22, True),
            ("Caption", "宋体", 10.5, False),
        ]
        for style_name, east_asia_font, size, bold in style_config:
            set_style_fonts(self.document.styles[style_name], east_asia_font=east_asia_font, ascii_font="Times New Roman", size=size, bold=bold)

    def _configure_sections(self) -> None:
        for section in self.document.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            section.header_distance = Cm(1.5)
            section.footer_distance = Cm(1.75)

    def _add_page_number(self, paragraph) -> None:
        add_field(paragraph, "PAGE", "")
        if paragraph.runs:
            set_run_fonts(paragraph.runs[-1], east_asia_font="Times New Roman", ascii_font="Times New Roman", size=10.5)

    def _body_section(self):
        section = self.document.add_section(WD_SECTION_START.NEW_PAGE)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(1.75)
        footer_p = section.footer.paragraphs[0]
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._add_page_number(footer_p)
        return section

    def cover_line(self, text: str, size: float = 14, bold: bool = False, center: bool = True, space_after: float = 0) -> None:
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(space_after)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(24)
        run = p.add_run(text)
        east_asia_font = "宋体" if any("\u4e00" <= ch <= "\u9fff" for ch in text) else "Times New Roman"
        set_run_fonts(run, east_asia_font=east_asia_font, ascii_font="Times New Roman", size=size, bold=bold)

    def front_heading(self, text: str, english: bool = False) -> None:
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(12)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(24)
        run = p.add_run(text)
        set_run_fonts(run, east_asia_font="Times New Roman" if english else "黑体", ascii_font="Times New Roman", size=16, bold=True)

    def chapter(self, text: str) -> None:
        p = self.document.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(12)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(24)
        set_run_fonts(p.add_run(text), east_asia_font="黑体", ascii_font="Times New Roman", size=16, bold=True)

    def section(self, text: str) -> None:
        p = self.document.add_paragraph(style="Heading 2")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(22)
        set_run_fonts(p.add_run(text), east_asia_font="黑体", ascii_font="Times New Roman", size=14, bold=True)

    def subsection(self, text: str) -> None:
        p = self.document.add_paragraph(style="Heading 3")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        set_run_fonts(p.add_run(text), east_asia_font="黑体", ascii_font="Times New Roman", size=12, bold=True)

    def body(self, text: str, english: bool = False, center: bool = False, indent: bool = True) -> None:
        p = self.document.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else (WD_ALIGN_PARAGRAPH.JUSTIFY if not english else WD_ALIGN_PARAGRAPH.LEFT)
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0.85) if indent else Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        east_asia_font = "Times New Roman" if english else "宋体"
        set_run_fonts(p.add_run(text), east_asia_font=east_asia_font, ascii_font="Times New Roman", size=12)

    def add_keywords(self, label: str, values: list[str] | str, english: bool = False) -> None:
        text = values if isinstance(values, str) else ("; ".join(values) if english else "；".join(values))
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(20)
        r1 = p.add_run(label)
        r2 = p.add_run(text)
        if english:
            set_run_fonts(r1, east_asia_font="Times New Roman", ascii_font="Times New Roman", size=12, bold=True)
            set_run_fonts(r2, east_asia_font="Times New Roman", ascii_font="Times New Roman", size=12)
        else:
            set_run_fonts(r1, east_asia_font="黑体", ascii_font="Times New Roman", size=12, bold=True)
            set_run_fonts(r2, east_asia_font="宋体", ascii_font="Times New Roman", size=12)

    def add_table_caption(self, text: str) -> None:
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(4)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(18)
        set_run_fonts(p.add_run(text), east_asia_font="宋体", ascii_font="Times New Roman", size=10.5, bold=True)

    def add_figure_caption(self, text: str) -> None:
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(3)
        pf.space_after = Pt(6)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(18)
        set_run_fonts(p.add_run(text), east_asia_font="宋体", ascii_font="Times New Roman", size=10.5)

    def set_cell_text(self, cell, text: str, bold: bool = False, center: bool = False) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(18)
        set_run_fonts(p.add_run(text), east_asia_font="宋体", ascii_font="Times New Roman", size=10.5, bold=bold)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def add_table(self, caption: str, rows: list[list[str]]) -> None:
        if not rows or not rows[0]:
            raise ValueError("Table rows must be a non-empty 2D array.")
        width = len(rows[0])
        if any(len(row) != width for row in rows):
            raise ValueError("All table rows must have the same width.")
        self.add_table_caption(caption)
        table = self.document.add_table(rows=len(rows), cols=width)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, row in enumerate(rows):
            for c_idx, value in enumerate(row):
                self.set_cell_text(table.cell(r_idx, c_idx), str(value), bold=(r_idx == 0), center=(r_idx == 0))
        self.document.add_paragraph()

    def resolve_path(self, raw_path: str) -> Path:
        path = Path(raw_path)
        if path.is_absolute():
            return path
        manifest_relative = (self.manifest_path.parent / path).resolve()
        if manifest_relative.exists():
            return manifest_relative
        if self.workspace is not None:
            workspace_relative = (self.workspace / path).resolve()
            if workspace_relative.exists():
                return workspace_relative
        return manifest_relative

    def add_figure(self, path_str: str, caption: str, max_width_cm: float, max_height_cm: float) -> None:
        image_path = self.resolve_path(path_str)
        if not image_path.exists():
            raise FileNotFoundError(f"Figure not found: {path_str} -> {image_path}")
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        width, height = fit_image(image_path, max_width_cm=max_width_cm, max_height_cm=max_height_cm)
        run = p.add_run()
        if height is not None:
            run.add_picture(str(image_path), width=width, height=height)
        else:
            run.add_picture(str(image_path), width=width)
        self.add_figure_caption(caption)

    def add_reference(self, text: str) -> None:
        p = self.document.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.left_indent = Cm(0.74)
        pf.first_line_indent = Cm(-0.74)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(18)
        set_run_fonts(p.add_run(text), east_asia_font="宋体", ascii_font="Times New Roman", size=10.5)

    def render_cover_and_front_matter(self) -> None:
        metadata = self.manifest.get("metadata", {})
        for _ in range(3):
            self.document.add_paragraph()
        self.cover_line(metadata.get("document_label", "毕业论文"), size=20, bold=True, space_after=6)
        self.cover_line(metadata["title_cn"], size=22, bold=True, space_after=10)
        if metadata.get("title_en"):
            self.cover_line(metadata["title_en"], size=14, space_after=20)
        for _ in range(3):
            self.document.add_paragraph()
        for label, key in [
            ("学生姓名：", "student_name"),
            ("学    号：", "student_id"),
            ("专    业：", "major"),
            ("学    院：", "college"),
            ("指导教师：", "advisor"),
        ]:
            self.cover_line(f"{label}{metadata.get(key, '________________')}", size=14, center=True, space_after=4)
        self.cover_line(f"完成日期：{metadata.get('date_cn', '________________')}", size=14, center=True, space_after=0)

        self._body_section()

        abstract_cn = metadata.get("abstract_cn") or []
        if abstract_cn:
            self.front_heading("摘要")
            for paragraph in abstract_cn:
                self.body(str(paragraph))
            if metadata.get("keywords_cn"):
                self.add_keywords("关键词：", metadata["keywords_cn"])
            self.document.add_page_break()

        abstract_en = metadata.get("abstract_en") or []
        if abstract_en:
            self.front_heading("Abstract", english=True)
            for paragraph in abstract_en:
                self.body(str(paragraph), english=True)
            if metadata.get("keywords_en"):
                self.add_keywords("Keywords: ", metadata["keywords_en"], english=True)
            self.document.add_page_break()

        if metadata.get("include_toc", True):
            self.front_heading("目录")
            toc_p = self.document.add_paragraph()
            toc_p.paragraph_format.space_before = Pt(0)
            toc_p.paragraph_format.space_after = Pt(0)
            toc_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            toc_p.paragraph_format.line_spacing = Pt(20)
            add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u', "目录将在打开文档后自动更新")
            self.document.add_page_break()

    def render_blocks(self) -> None:
        blocks = self.manifest.get("blocks") or self.manifest.get("content_blocks")
        if not isinstance(blocks, list):
            raise ValueError("Manifest must contain a top-level 'blocks' array.")

        for index, block in enumerate(blocks, start=1):
            if not isinstance(block, dict):
                raise ValueError(f"Block #{index} must be an object.")
            block_type = block.get("type")
            if block_type == "chapter":
                self.chapter(block["title"])
            elif block_type == "section":
                self.section(block["title"])
            elif block_type == "subsection":
                self.subsection(block["title"])
            elif block_type == "paragraph":
                self.body(
                    block["text"],
                    english=bool(block.get("english", False)),
                    center=bool(block.get("center", False)),
                    indent=block.get("indent", True),
                )
            elif block_type == "figure":
                self.add_figure(
                    block["path"],
                    block["caption"],
                    float(block.get("max_width_cm", 15.2)),
                    float(block.get("max_height_cm", 18.5)),
                )
            elif block_type == "table":
                self.add_table(block["caption"], block["rows"])
            elif block_type == "page_break":
                self.document.add_page_break()
            elif block_type == "references":
                self.chapter(block.get("title", "参考文献"))
                for item in block.get("items", []):
                    self.add_reference(str(item))
            else:
                raise ValueError(f"Unsupported block type at index {index}: {block_type}")

    def build(self) -> None:
        metadata = self.manifest.get("metadata", {})
        title_cn = metadata.get("title_cn")
        if not title_cn:
            raise ValueError("Manifest metadata.title_cn is required.")
        self.document.core_properties.title = title_cn
        self.document.core_properties.subject = metadata.get("subject", "Software thesis document")
        self.document.core_properties.category = metadata.get("category", "thesis")
        self.document.core_properties.comments = metadata.get("comments", "Generated from a structured thesis manifest.")
        self.render_cover_and_front_matter()
        self.render_blocks()
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(self.output_path))


def load_manifest(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as fh:
        data = json.load(fh)
    if not isinstance(data, dict):
        raise ValueError("Manifest root must be a JSON object.")
    return data


def main() -> None:
    args = parse_args()
    manifest_path = Path(args.manifest).resolve()
    output_path = Path(args.output).resolve()
    workspace = Path(args.workspace).resolve() if args.workspace else None
    manifest = load_manifest(manifest_path)
    builder = ThesisDocBuilder(manifest, manifest_path=manifest_path, output_path=output_path, workspace=workspace)
    builder.build()
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()

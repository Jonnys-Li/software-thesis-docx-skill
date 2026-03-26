from __future__ import annotations

import argparse
import json
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from PIL import Image


NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


@dataclass(frozen=True)
class MappingRule:
    caption: str
    image_path: Path
    fit_mode: str


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Replace inline DOCX images by matching the following caption paragraph.")
    parser.add_argument("--input", required=True, help="Input DOCX path.")
    parser.add_argument("--output", required=True, help="Output DOCX path.")
    parser.add_argument("--mapping", required=True, help="JSON file containing caption -> image replacements.")
    return parser.parse_args()


def load_mapping(mapping_path: Path) -> list[MappingRule]:
    with mapping_path.open("r", encoding="utf-8") as fh:
        data = json.load(fh)
    items: Any
    if isinstance(data, list):
        items = data
    elif isinstance(data, dict) and isinstance(data.get("items"), list):
        items = data["items"]
    else:
        raise ValueError("Mapping file must be a list or an object with an 'items' array.")

    rules: list[MappingRule] = []
    for item in items:
        caption = item["caption"]
        fit_mode = item["fit_mode"]
        if fit_mode not in {"original_box", "page_width"}:
            raise ValueError(f"Unsupported fit_mode '{fit_mode}' for caption '{caption}'.")
        image_path = Path(item["image_path"])
        if not image_path.is_absolute():
            image_path = (mapping_path.parent / image_path).resolve()
        rules.append(MappingRule(caption=caption, image_path=image_path, fit_mode=fit_mode))
    return rules


def iter_figure_slots(document: Document, target_captions: set[str]) -> dict[str, object]:
    slots: dict[str, object] = {}
    for idx, paragraph in enumerate(document.paragraphs):
        if paragraph.text not in target_captions:
            continue
        if paragraph.text in slots:
            raise ValueError(f"Duplicate figure caption found: {paragraph.text}")
        if idx == 0:
            raise ValueError(f"Caption has no preceding paragraph: {paragraph.text}")
        image_paragraph = document.paragraphs[idx - 1]
        has_drawing = bool(image_paragraph._element.findall(".//w:drawing", NS))
        if not has_drawing:
            raise ValueError(f"Caption is not preceded by an image paragraph: {paragraph.text}")
        slots[paragraph.text] = image_paragraph
    return slots


def get_image_geometry(image_path: Path) -> tuple[int, int]:
    with Image.open(image_path) as img:
        width_px, height_px = img.size
    if width_px <= 0 or height_px <= 0:
        raise ValueError(f"Invalid image dimensions: {image_path}")
    return width_px, height_px


def fit_into_bounds(image_path: Path, max_width_emu: int, max_height_emu: int) -> tuple[int, int]:
    width_px, height_px = get_image_geometry(image_path)
    aspect_ratio = width_px / height_px
    width_emu = int(max_width_emu)
    height_emu = int(round(width_emu / aspect_ratio))
    if height_emu > max_height_emu:
        height_emu = int(max_height_emu)
        width_emu = int(round(height_emu * aspect_ratio))
    return width_emu, height_emu


def set_picture_extent(image_paragraph, width_emu: int, height_emu: int) -> None:
    extent = image_paragraph._element.find(".//wp:extent", NS)
    shape_extent = image_paragraph._element.find(".//pic:spPr/a:xfrm/a:ext", NS)
    if extent is None or shape_extent is None:
        raise ValueError("Unable to locate picture extent nodes.")
    extent.set("cx", str(width_emu))
    extent.set("cy", str(height_emu))
    shape_extent.set("cx", str(width_emu))
    shape_extent.set("cy", str(height_emu))


def get_current_extent(image_paragraph) -> tuple[int, int]:
    extent = image_paragraph._element.find(".//wp:extent", NS)
    if extent is None:
        raise ValueError("Unable to locate current picture extent.")
    return int(extent.get("cx")), int(extent.get("cy"))


def replace_image_blob(document: Document, image_paragraph, replacement_image: Path) -> None:
    if not replacement_image.exists():
        raise FileNotFoundError(f"Replacement image not found: {replacement_image}")
    blip = image_paragraph._element.find(".//a:blip", NS)
    if blip is None:
        raise ValueError(f"Unable to locate image relationship for {replacement_image.name}")
    rel_id = blip.get(qn("r:embed"))
    if rel_id is None:
        raise ValueError(f"Picture is missing embed relationship for {replacement_image.name}")
    image_part = document.part.related_parts[rel_id]
    image_part._blob = replacement_image.read_bytes()


def usable_page_width(document: Document) -> int:
    return min(section.page_width - section.left_margin - section.right_margin for section in document.sections)


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()
    mapping_path = Path(args.mapping).resolve()
    rules = load_mapping(mapping_path)

    shutil.copy2(input_path, output_path)
    document = Document(output_path)
    required_captions = {rule.caption for rule in rules}
    slots = iter_figure_slots(document, required_captions)
    missing = sorted(required_captions - slots.keys())
    if missing:
        raise ValueError(f"Missing figure caption(s) in document: {missing}")

    page_width_emu = usable_page_width(document)

    for rule in rules:
        image_paragraph = slots[rule.caption]
        old_width_emu, old_height_emu = get_current_extent(image_paragraph)
        replace_image_blob(document, image_paragraph, rule.image_path)

        if rule.fit_mode == "page_width":
            width_emu, height_emu = fit_into_bounds(rule.image_path, max_width_emu=page_width_emu, max_height_emu=old_height_emu)
        else:
            width_emu, height_emu = fit_into_bounds(rule.image_path, max_width_emu=old_width_emu, max_height_emu=old_height_emu)

        set_picture_extent(image_paragraph, width_emu, height_emu)

    document.save(output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
